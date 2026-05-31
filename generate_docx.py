from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
TEAL   = RGBColor(0x00, 0xBF, 0xA6)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF5, 0xF5, 0xF5)
MGREY  = RGBColor(0xCC, 0xCC, 0xCC)
BLACK  = RGBColor(0x00, 0x00, 0x00)
AMBER  = RGBColor(0xF5, 0xA6, 0x23)
GREEN  = RGBColor(0x22, 0xC5, 0x5E)
RED    = RGBColor(0xEF, 0x44, 0x44)

# ── Helper: set cell shading ──────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    kwargs[edge].get('sz','4'))
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: paragraph style shortcuts ────────────────────────────────────────
def add_heading(doc, text, level=1, color=DARK):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p

def add_para(doc, text='', bold=False, italic=False, size=10, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

# ── Helper: build a styled table ─────────────────────────────────────────────
def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, '1A1A2E')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            shade_cell(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)

    # column widths
    if col_widths:
        for ri2, row2 in enumerate(table.rows):
            for ci2, w in enumerate(col_widths):
                row2.cells[ci2].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
# Teal header bar (table trick)
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_cell = cover_tbl.rows[0].cells[0]
shade_cell(cover_cell, '00BFA6')
cover_cell.width = Inches(6.5)
cp = cover_cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run('MANSCO SPARE PARTS PORTAL')
cr.bold = True
cr.font.size = Pt(22)
cr.font.color.rgb = WHITE

doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('Milestone Delivery Plan & Project Schedule')
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = DARK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('Version 1.0 — Prepared for Customer Sign-Off')
sr.italic = True
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# Meta table
meta_tbl = doc.add_table(rows=5, cols=2)
meta_tbl.style = 'Table Grid'
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Document Reference', 'MANSCO-SPP-MDP-001'),
    ('Prepared by',        'SmartForce Technology Solutions'),
    ('Prepared for',       'MANSCO Egypt — Peugeot Spare Parts Division'),
    ('Document Date',      'May 20, 2026'),
    ('Document Status',    'Awaiting Customer Approval'),
]
for ri, (label, val) in enumerate(meta_data):
    lc = meta_tbl.rows[ri].cells[0]
    vc = meta_tbl.rows[ri].cells[1]
    shade_cell(lc, 'E8F8F6')
    shade_cell(vc, 'FFFFFF')
    lp = lc.paragraphs[0]; lp.add_run(label).bold = True
    lp.runs[0].font.size = Pt(10)
    vp = vc.paragraphs[0]; vp.add_run(val)
    vp.runs[0].font.size = Pt(10)

for row in meta_tbl.rows:
    row.cells[0].width = Inches(2.2)
    row.cells[1].width = Inches(4.3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Project Overview', level=1, color=TEAL)

add_heading(doc, 'Executive Summary', level=2, color=DARK)
add_para(doc,
    'The MANSCO Spare Parts Portal is an enterprise-grade, web-based self-service platform '
    'developed for MANSCO Egypt — the authorized distributor of Peugeot spare parts in the '
    'Egyptian market. The system replaces a fully manual, representative-mediated spare parts '
    'ordering process with a governed, digital self-service platform accessible to an authorized '
    'network of dealers and sub-dealers.',
    size=10)
doc.add_paragraph()
add_para(doc,
    'The portal covers the full spare parts lifecycle: dealer onboarding and approval, part search '
    'with availability-aware pricing, multi-type order placement (Daily, Air/DHL, Stock), order '
    'tracking, invoice management, back-order handling, financial account management, '
    'campaign/promotion management, and operational reporting.',
    size=10)
doc.add_paragraph()
add_para(doc,
    'The system integrates offline with SAP (MANSCO\'s ERP system) via a structured, batch-based '
    'CSV data exchange protocol. SAP remains the single source of truth for inventory, pricing, '
    'and business rules, while the portal acts as the dealer-facing interaction layer.',
    size=10)

add_heading(doc, 'Business Objectives Addressed', level=2, color=DARK)
biz_headers = ['Objective', 'System Response']
biz_rows = [
    ('Eliminate manual phone/email order processing',
     'Self-service order placement with real-time availability checks'),
    ('Enforce financial controls before order submission',
     'Credit limit, overdue balance, and financial status checks at order-time'),
    ('Enforce pricing rules (no price for unavailable items)',
     'Availability-aware pricing engine: prices hidden for out-of-stock items'),
    ('Streamline dealer onboarding',
     'Digital registration → email verification → admin approval workflow'),
    ('Enable marketing campaigns with automated discounts',
     'Campaign Manager module with eligibility rules and automatic discount application'),
    ('Provide management visibility into dealer activity',
     'Admin operational dashboard, inquiry reports, lost-sales reports'),
    ('Support bilingual Arabic/English interface',
     'Full RTL/LTR i18n framework with EN/AR translation coverage'),
]
make_table(doc, biz_headers, biz_rows, col_widths=[2.8, 3.7])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE & TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '2. Architecture & Tech Stack Summary', level=1, color=TEAL)

add_heading(doc, 'Technology Identification', level=2, color=DARK)
tech_headers = ['Layer', 'Technology', 'Version', 'Purpose']
tech_rows = [
    ('Frontend Framework', 'Next.js (App Router)', '16.2.6', 'Full-stack React SSR/SSG framework'),
    ('UI Runtime', 'React', '19.2.4', 'Component rendering engine'),
    ('Language', 'TypeScript', '^5', 'Type-safe development across all layers'),
    ('Styling', 'Tailwind CSS', '^4', 'Utility-first CSS with dark-theme design system'),
    ('Component Library', 'shadcn/ui + Base UI', '^4.7.0 / ^1.4.1', 'Accessible, composable UI primitives'),
    ('Data Visualization', 'Recharts', '^3.8.1', 'Dashboard KPI charts and analytics'),
    ('Database', 'Supabase PostgreSQL', 'Cloud-hosted', 'Primary relational database'),
    ('Authentication', 'Supabase Auth', '^2.106.0', 'Email/password auth, session management, JWTs'),
    ('Storage', 'Supabase Storage', '—', 'Dealer document uploads'),
    ('API Layer', 'Next.js Route Handlers', '—', 'RESTful API endpoints (server-side)'),
    ('Input Validation', 'Zod', '^4.4.3', 'Schema validation on all API inputs'),
    ('Internationalization', 'next-intl + custom i18n', '^4.12.0', 'EN/AR bilingual support with RTL layout'),
]
make_table(doc, tech_headers, tech_rows, col_widths=[1.5, 1.7, 1.2, 2.1])
doc.add_paragraph()

add_heading(doc, 'Database Schema Summary', level=2, color=DARK)
db_headers = ['Table', 'Status', 'Purpose']
db_rows = [
    ('dealer_registrations', '✅ Deployed', 'Dealer onboarding applications and document tracking'),
    ('dealers',              '✅ Deployed', 'Approved dealer master records with financial data'),
    ('campaigns',            '✅ Deployed', 'Campaign definitions (type, dates, audience, status)'),
    ('campaign_items',       '✅ Deployed', 'Parts assigned to campaigns with discount configuration'),
    ('campaign_orders',      '✅ Deployed', 'Campaign participation tracking per order'),
    ('campaign_audit_log',   '✅ Deployed', 'Full audit trail of campaign lifecycle changes'),
    ('orders',               '🔲 Pending',  'Order headers (type, status, dealer, totals)'),
    ('order_lines',          '🔲 Pending',  'Individual part lines per order'),
    ('order_approvals',      '🔲 Pending',  'Admin review actions on orders'),
    ('invoices',             '🔲 Pending',  'Invoice records linked to fulfilled orders'),
    ('shipments',            '🔲 Pending',  'Carrier/tracking data per shipment'),
    ('back_orders',          '🔲 Pending',  'Back-order lines with ETA tracking'),
    ('inquiries',            '🔲 Pending',  'Part search and order attempt logging'),
    ('lost_sales',           '🔲 Pending',  'Lost sale candidates for demand planning'),
    ('stock_availability',   '🔲 Pending',  'SAP-synced stock data per part/location'),
    ('price_lists',          '🔲 Pending',  'SAP-synced price list headers'),
    ('price_list_items',     '🔲 Pending',  'Part-level pricing with discounts'),
    ('sync_logs',            '🔲 Pending',  'SAP CSV import/export audit log'),
    ('dealer_targets',       '🔲 Pending',  'KPI targets vs. achievement per dealer'),
]
make_table(doc, db_headers, db_rows, col_widths=[1.8, 1.0, 3.7])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM MODULES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '3. System Modules', level=1, color=TEAL)

modules = [
    {
        'num': 1, 'title': 'Authentication & Dealer Onboarding',
        'status': '✅ Delivered',
        'items': [
            'Supabase-backed email/password authentication',
            'Dealer self-registration form with Zod validation',
            'Document upload to Supabase Storage (Tax Card, Trade License, CR Document)',
            'Email verification flow with post-verification redirect',
            'Pending approval holding screen',
            'Admin registration review queue, document preview, approve/reject with notes',
            'Auth middleware (route-level access control by role and registration_status)',
            'Demo admin bypass for development/testing',
        ]
    },
    {
        'num': 2, 'title': 'Dealer Management (Admin)',
        'status': '✅ Delivered',
        'items': [
            'Full dealer listing with real-time Supabase data',
            'Search and filter by company, code, email, address',
            'Dealer preview side panel (all profile and financial details)',
            'Dealer edit modal (code, financials, status, type, contact details)',
            'PUT /api/dealers/[id] API with field whitelisting',
            'Financial status management (Active / Warning / Blocked)',
            'Credit limit and overdue balance management',
        ]
    },
    {
        'num': 3, 'title': 'Campaign Manager (Admin)',
        'status': '✅ Delivered',
        'items': [
            'Campaign lifecycle: Draft → Active → Paused → Completed → Archived',
            'Campaign creation wizard (name, type, dates, audience targeting, eligibility rules)',
            'Campaign item management (parts, discount type %, EGP, min order quantity)',
            'Eligibility rules engine (min credit limit, min target %, financial status, order type)',
            'Campaign performance analytics (orders, dealers, quantities, discounts)',
            'Dealer participation table per campaign',
            'Full audit trail per campaign',
            'Duplicate campaign (clone as new draft)',
            '6 REST API endpoints covering all CRUD + status transitions + performance',
        ]
    },
    {
        'num': 4, 'title': 'Parts Catalog & Search',
        'status': '🟡 UI Delivered / Backend Integration Pending',
        'items': [
            'Part search with category and model filters',
            'Availability-aware display (4 states: Available, Partially Available, Out with ETA, Out no ETA)',
            'Critical business rule enforced: prices hidden for out-of-stock items',
            'Multi-language part names and categories (EN/AR)',
            'Supports Peugeot model range: 301, 208, 508, 2008, 3008, 5008, Partner, Expert, Boxer',
            '10 spare part categories (Brakes, Engine, Electrical, Body, Filters, Cooling, etc.)',
            'Currently using mock data; requires SAP CSV sync integration to go live',
        ]
    },
    {
        'num': 5, 'title': 'Order Management',
        'status': '🟡 UI Delivered / Backend Persistence Pending',
        'items': [
            'New order flow: part search → cart → type selection → submission',
            'Three order types: Daily, Air/DHL, Stock (each with distinct ETA logic)',
            'Cart management: add/remove/quantity adjustment',
            'Financial block check before order submission (credit limit, overdue balance)',
            'Order history list with status filtering',
            'Order detail view with line-level status and tracking',
            'Admin order review queue (approve, partial approve, reject)',
            'Status lifecycle: Submitted → Under Review → Approved/Rejected → Done → Invoiced → Shipped → Delivered',
        ]
    },
    {
        'num': 6, 'title': 'Fulfillment Tracking (Invoices, Shipments, Back-Orders)',
        'status': '🔲 Pending',
        'items': [
            'Invoice listing and detail view (linked to fulfilled orders)',
            'Shipment tracking with carrier/AWB/DHL reference',
            'Back-order management with ETA tracking and recalculation',
            'Partial fulfillment workflow (confirmed lines vs. back-order lines)',
        ]
    },
    {
        'num': 7, 'title': 'Financial Dashboard & KPIs (Dealer-Facing)',
        'status': '🟡 UI Delivered / Backend Integration Pending',
        'items': [
            'KPI cards: Credit Limit Utilization, Overdue Balance, Target Achievement',
            'Target vs. achievement progress visualization',
            'Recent orders widget',
            'Campaign banners and discounted items',
            'Financial summary with credit/overdue status indicators',
            'Currently using mock data; requires live Supabase data binding',
        ]
    },
    {
        'num': 8, 'title': 'Admin Operational Dashboard & Reporting',
        'status': '🟡 UI Delivered / Backend Integration Pending',
        'items': [
            'Admin overview: pending approvals, today\'s orders, revenue, active campaigns',
            'Inquiry report (all part search activity by dealer)',
            'Lost-sales report (unfulfilled inquiries for demand planning)',
            'Currently using mock data; requires database and sync integration',
        ]
    },
    {
        'num': 9, 'title': 'SAP CSV Integration Engine',
        'status': '🔲 Not Started',
        'items': [
            'Scheduled CSV import: stock availability, pricing, invoices from SAP',
            'Scheduled CSV export: new orders to SAP',
            'Papaparse-based CSV parser with schema validation',
            'SyncLog table tracking per-run record counts and errors',
            'Stale-data detection and UI warning badges',
            'Archive/staging workflow for processed files in Supabase Storage',
        ]
    },
]

status_color = {
    '✅ Delivered': '00BFA6',
    '🟡 UI Delivered / Backend Integration Pending': 'F5A623',
    '🟡 UI Delivered / Backend Persistence Pending': 'F5A623',
    '🔲 Pending': 'CCCCCC',
    '🔲 Not Started': 'CCCCCC',
}

for mod in modules:
    p = doc.add_paragraph()
    r1 = p.add_run(f'Module {mod["num"]} — {mod["title"]}')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK

    # status badge inline
    r2 = p.add_run(f'   {mod["status"]}')
    r2.font.size = Pt(9)
    sc = status_color.get(mod['status'], 'CCCCCC')
    r2.font.color.rgb = RGBColor(int(sc[0:2],16), int(sc[2:4],16), int(sc[4:6],16))
    r2.bold = True

    for item in mod['items']:
        add_bullet(doc, item)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CURRENT DELIVERY STATUS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '4. Current Delivery Status', level=1, color=TEAL)

status_headers = ['Module', 'Phase', 'Status', '% Complete']
status_rows = [
    ('Auth & Dealer Onboarding',         'Phase 1',   '✅ Complete',   '100%'),
    ('Dealer Management (Admin)',         'Phase 1',   '✅ Complete',   '100%'),
    ('Campaign Manager',                  'Phase 2',   '✅ Complete',   '100%'),
    ('Core UI Component Library',         'Phase 1',   '✅ Complete',   '100%'),
    ('Bilingual EN/AR Framework',         'Phase 1',   '✅ Complete',   '100%'),
    ('Parts Catalog & Search (UI)',       'Phase 2',   '🟡 In Progress','70%'),
    ('Order Management (UI)',             'Phase 2',   '🟡 In Progress','65%'),
    ('Financial Dashboard (UI)',          'Phase 2',   '🟡 In Progress','60%'),
    ('Admin Dashboard & Reports (UI)',    'Phase 2',   '🟡 In Progress','55%'),
    ('Database — Full Schema',            'Phase 2–3', '🔲 Partial',    '30%'),
    ('Order Backend APIs',                'Phase 3',   '🔲 Pending',    '0%'),
    ('Fulfillment (Invoices/Shipments)',  'Phase 3',   '🔲 Pending',    '0%'),
    ('SAP CSV Sync Engine',               'Phase 3',   '🔲 Pending',    '0%'),
    ('Rule Engine (Order Validation)',    'Phase 3',   '🔲 Pending',    '0%'),
    ('UAT & Production Deployment',       'Phase 4',   '🔲 Pending',    '0%'),
]
make_table(doc, status_headers, status_rows, col_widths=[2.5, 1.0, 1.8, 1.2])
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Overall Project Completion: ~38%')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = TEAL

# ══════════════════════════════════════════════════════════════════════════════
# 5. MILESTONE DELIVERY PLAN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '5. Milestone Delivery Plan', level=1, color=TEAL)

add_para(doc,
    'Note on Duration Estimates: All estimates assume a single experienced full-stack development '
    'team (2 developers + 1 QA). Durations are based on the complexity observed in the existing '
    'codebase and include development, integration testing, and code review. Client UAT time is '
    'tracked separately.',
    italic=True, size=9, color=RGBColor(0x55,0x55,0x55))
doc.add_paragraph()

ms_headers = ['#', 'Milestone Phase', 'Key Deliverables', 'Duration', 'Start', 'Dependencies']
ms_rows = [
    ('M1', 'Foundation & Core Architecture (Delivered)',
     'Auth system, dealer registration, middleware RBAC, UI component library, bilingual framework, Supabase infrastructure',
     '4 weeks', '✅ Complete', 'Supabase project provisioning, domain setup'),
    ('M2', 'Dealer & Campaign Management (Delivered)',
     'Dealer management with edit/preview, campaign full lifecycle, campaign DB schema deployed',
     '3 weeks', '✅ Complete', 'M1 complete, Supabase migrations'),
    ('M3', 'Parts Catalog & Live Data Binding',
     'Connect parts search to live data, stock availability API, pricing API, availability state engine, remove mock data',
     '2 weeks', 'Week 1', 'M1 complete, SAP CSV format confirmed'),
    ('M4', 'Order Management Backend',
     'Full order DB schema, order CRUD APIs, validation rule engine, admin approval queue, partial fulfillment',
     '4 weeks', 'Week 3', 'M3 complete, DB schema finalized'),
    ('M5', 'Fulfillment & Financial Tracking',
     'Invoice model, shipment tracking, back-order lifecycle, ETA recalculation, dealer financial dashboard wired to live data',
     '3 weeks', 'Week 7', 'M4 complete'),
    ('M6', 'SAP CSV Integration Engine',
     'CSV import engine, CSV export engine, cron scheduler, SyncLog tracking, stale-data warnings',
     '3 weeks', 'Week 10', 'M4 complete, SAP CSV schemas confirmed'),
    ('M7', 'Reporting, Analytics & Admin Dashboard',
     'Inquiry logging, lost-sales logging, admin dashboard live data, inquiry and lost-sales reports, CSV/Excel export',
     '2 weeks', 'Week 13', 'M5, M6 complete'),
    ('M8', 'Hardening, Security & Performance',
     'Security audit, RLS policy review, query optimization, load testing, error monitoring',
     '2 weeks', 'Week 15', 'M7 complete'),
    ('M9', 'UAT & Customer Acceptance',
     'UAT with MANSCO team, bug triage and resolution, dealer pilot onboarding, sign-off on all business rules',
     '2 weeks', 'Week 17', 'M8 complete, UAT test plan agreed'),
    ('M10', 'Production Deployment & Go-Live',
     'Production environment setup, CI/CD pipeline, DNS/domain cutover, monitoring & alerting, go-live support',
     '1 week', 'Week 19', 'M9 sign-off, production infrastructure ready'),
]
make_table(doc, ms_headers, ms_rows, col_widths=[0.4, 1.6, 2.5, 0.7, 0.6, 1.2])
doc.add_paragraph()

# Gantt (text representation)
add_heading(doc, 'Schedule Overview', level=2, color=DARK)
gantt = doc.add_paragraph()
gantt.add_run(
    'Week:   1  2  3  4  5  6  7  8  9  10  11  12  13  14  15  16  17  18  19\n'
    'M1–M2:  [DELIVERED ────────────────]\n'
    'M3:                                  [──────]\n'
    'M4:                                          [──────────────]\n'
    'M5:                                                          [────────]\n'
    'M6:                                                          [────────]\n'
    'M7:                                                                    [──────]\n'
    'M8:                                                                            [──────]\n'
    'M9:                                                                                    [──────]\n'
    'M10:                                                                                           [───]'
).font.name = 'Courier New'
gantt.runs[0].font.size = Pt(8)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Total Remaining Duration: ~19 weeks from project resumption date')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
p2 = doc.add_paragraph()
r2 = p2.add_run('Total Project Duration (including delivered phases): ~24 weeks')
r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = DARK

# ══════════════════════════════════════════════════════════════════════════════
# 6. DETAILED PHASE BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '6. Detailed Phase Breakdown', level=1, color=TEAL)

phases = [
    {
        'title': 'M3 — Parts Catalog & Live Data Binding (2 weeks)',
        'headers': ['Task', 'Effort', 'Owner'],
        'rows': [
            ('Deploy stock_availability, price_lists, price_list_items DB tables', '0.5d', 'Backend'),
            ('Build GET /api/parts/search with pagination and filters', '1d', 'Backend'),
            ('Build GET /api/parts/[partNumber]/availability', '0.5d', 'Backend'),
            ('Implement 4-state availability logic', '1d', 'Backend'),
            ('Enforce no-price rule for out-of-stock states', '0.5d', 'Backend'),
            ('Wire Parts Catalog UI to live API (remove mock data)', '1d', 'Frontend'),
            ('Wire Dashboard KPI cards to Supabase (remove mock data)', '1d', 'Frontend'),
            ('Integration tests for pricing rules', '1d', 'QA'),
        ],
        'widths': [3.8, 0.7, 1.0],
    },
    {
        'title': 'M4 — Order Management Backend (4 weeks)',
        'headers': ['Task', 'Effort', 'Owner'],
        'rows': [
            ('Deploy orders, order_lines, order_approvals DB schema', '1d', 'Backend'),
            ('Build order validation rule engine (5-step chain)', '4d', 'Backend'),
            ('Build POST /api/orders (create order)', '2d', 'Backend'),
            ('Build GET /api/orders and GET /api/orders/[id]', '1d', 'Backend'),
            ('Build POST /api/orders/[id]/review (admin approve/reject)', '1d', 'Backend'),
            ('Wire cart-to-order submission flow in frontend', '2d', 'Frontend'),
            ('Wire order history and order detail pages to live API', '2d', 'Frontend'),
            ('Wire admin approval queue to live API', '1d', 'Frontend'),
            ('Unit tests: rule engine (eligibility, credit, quota, pricing, financial)', '3d', 'QA'),
            ('Integration tests: full order lifecycle', '2d', 'QA'),
        ],
        'widths': [3.8, 0.7, 1.0],
    },
    {
        'title': 'M5 — Fulfillment & Financial Tracking (3 weeks)',
        'headers': ['Task', 'Effort', 'Owner'],
        'rows': [
            ('Deploy invoices, shipments, back_orders DB schema', '1d', 'Backend'),
            ('Build invoice APIs (list, detail)', '1d', 'Backend'),
            ('Build shipment tracking API', '1d', 'Backend'),
            ('Build back-order lifecycle API (create, update ETA, resolve)', '2d', 'Backend'),
            ('Partial fulfillment split logic', '2d', 'Backend'),
            ('Wire Invoices page to live API', '1d', 'Frontend'),
            ('Wire Back-Orders page to live API', '1d', 'Frontend'),
            ('Dealer financial dashboard — live data binding', '1.5d', 'Frontend'),
            ('Integration tests', '2d', 'QA'),
        ],
        'widths': [3.8, 0.7, 1.0],
    },
    {
        'title': 'M6 — SAP CSV Integration Engine (3 weeks)',
        'headers': ['Task', 'Effort', 'Owner'],
        'rows': [
            ('Define and validate CSV schemas with MANSCO/SAP team', '2d', 'Analyst'),
            ('Build CSV import engine (stock availability)', '2d', 'Backend'),
            ('Build CSV import engine (pricing data)', '1.5d', 'Backend'),
            ('Build CSV import engine (invoices from SAP)', '1d', 'Backend'),
            ('Build CSV export engine (orders to SAP)', '1.5d', 'Backend'),
            ('Build cron scheduler (configurable interval)', '1d', 'Backend'),
            ('SyncLog tracking and error reporting', '1d', 'Backend'),
            ('Supabase Storage staging/archive workflow', '1d', 'Backend'),
            ('Stale-data UI warnings (configurable threshold)', '0.5d', 'Frontend'),
            ('End-to-end sync testing with sample SAP files', '3d', 'QA'),
        ],
        'widths': [3.8, 0.7, 1.0],
    },
]

for phase in phases:
    add_heading(doc, phase['title'], level=2, color=DARK)
    make_table(doc, phase['headers'], phase['rows'], col_widths=phase['widths'])
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. ASSUMPTIONS & SCOPE LIMITS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '7. Assumptions & Scope Limits', level=1, color=TEAL)

add_heading(doc, '7.1 Technical Assumptions', level=2, color=DARK)
assumptions = [
    ('A1', 'The Supabase project (Cloud plan) remains provisioned and accessible. Any capacity upgrades for production load are the customer\'s responsibility.'),
    ('A2', 'SAP CSV file formats (column definitions, delimiters, encoding, file naming) will be provided and signed off by the MANSCO/SAP team before M6 begins.'),
    ('A3', 'MANSCO will provide test SAP CSV files (realistic sample data) no later than 2 weeks before M6 start.'),
    ('A4', 'SAP integration is batch-based (file-based, not real-time API). Real-time SAP API integration is explicitly out of scope.'),
    ('A5', 'Email delivery relies on Supabase\'s built-in SMTP or a configured relay. Production SMTP configuration is the customer\'s responsibility.'),
    ('A6', 'Production hosting infrastructure (Vercel, AWS, or equivalent) and domain/SSL will be provisioned by the customer.'),
    ('A7', 'The UI component library and all open-source dependencies are used under their respective licenses. No proprietary UI components are included.'),
    ('A8', 'Arabic translation content is included. Any business-specific Arabic terminology changes after delivery constitute a change request.'),
]
make_table(doc, ['#', 'Assumption'], assumptions, col_widths=[0.4, 6.1])
doc.add_paragraph()

add_heading(doc, '7.2 Scope Exclusions (Explicitly Out of Scope)', level=2, color=DARK)
exclusions = [
    ('E1', 'Real-time SAP integration — All SAP data exchange is batch CSV-based. Webhook, API, or event-driven SAP integration is not in scope.'),
    ('E2', 'Mobile native applications — The portal is a responsive web application. Native iOS/Android apps are not in scope.'),
    ('E3', 'Payment processing — No payment gateway integration is included. Financial account management is informational only.'),
    ('E4', 'ERP master data management — Creation or modification of dealers, parts, or price lists in SAP is not in scope.'),
    ('E5', 'Third-party logistics integrations — Direct API integrations with DHL, FedEx, or carriers are not in scope.'),
    ('E6', 'Custom BI / data warehouse — Advanced analytics beyond the defined reports are not in scope.'),
    ('E7', 'Load balancing and auto-scaling infrastructure — Beyond standard Vercel/Next.js deployment is not in scope.'),
    ('E8', 'User training and change management — Training materials and onboarding sessions are not included in technical delivery.'),
]
make_table(doc, ['#', 'Exclusion'], exclusions, col_widths=[0.4, 6.1])
doc.add_paragraph()

add_heading(doc, '7.3 Change Request Policy', level=2, color=DARK)
add_para(doc,
    'Any work outside the scope defined in this document will be handled via a formal Change Request (CR) process. '
    'A CR will include: description of the change and business justification, estimated additional effort and cost, '
    'impact on the delivery schedule, and mutual sign-off before implementation begins.',
    size=10)

# ══════════════════════════════════════════════════════════════════════════════
# 8. RISK REGISTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '8. Risk Register', level=1, color=TEAL)

risk_headers = ['#', 'Risk', 'Probability', 'Impact', 'Mitigation']
risk_rows = [
    ('R1', 'SAP CSV schemas not available before M6 start, causing delay',
     'Medium', 'High',
     'Confirm CSV formats as a milestone gate. Start M6 only after sample files received.'),
    ('R2', 'MANSCO dealer network volume exceeds Supabase free tier limits during UAT',
     'Low', 'Medium',
     'Monitor row counts and storage. Plan Supabase Pro upgrade before UAT.'),
    ('R3', 'Order rule engine edge cases discovered during UAT',
     'Medium', 'Medium',
     'Comprehensive unit tests in M4. Allocate buffer days in UAT phase for rule fixes.'),
    ('R4', 'Arabic RTL layout issues with complex table/form layouts',
     'Low', 'Low',
     'RTL testing included in each phase. Bilingual framework already implemented.'),
    ('R5', 'Supabase service outage affecting development or production',
     'Low', 'High',
     'Development continues with local Supabase CLI fallback. Production SLA monitored.'),
    ('R6', 'Customer feedback during UAT requires significant rework',
     'Medium', 'High',
     'Demo/review checkpoints at end of each milestone to surface issues early.'),
]
make_table(doc, risk_headers, risk_rows, col_widths=[0.4, 2.0, 0.9, 0.7, 2.5])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 9. CUSTOMER SIGN-OFF
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '9. Customer Sign-Off', level=1, color=TEAL)

add_heading(doc, 'Document Review & Acceptance', level=2, color=DARK)
add_para(doc,
    'By signing below, the authorized representatives of MANSCO Egypt and SmartForce Technology Solutions '
    'confirm that:', size=10)
items_sign = [
    'The scope of work described in this Milestone Delivery Plan is agreed upon and understood by both parties.',
    'The milestone schedule and duration estimates are accepted as the project baseline.',
    'Any work outside the defined scope will be processed through the formal Change Request procedure.',
    'Both parties commit to the responsibilities and assumptions outlined in Section 7.',
]
for item in items_sign:
    add_bullet(doc, item)
doc.add_paragraph()

# MANSCO sign block
add_heading(doc, 'MANSCO Egypt — Customer Authorization', level=2, color=DARK)
mansco_headers = ['Field', 'Details']
mansco_rows = [
    ('Company Name',              'MANSCO Egypt'),
    ('Authorized Representative', '________________________________'),
    ('Title / Position',          '________________________________'),
    ('Signature',                 '________________________________'),
    ('Date',                      '________________________________'),
]
make_table(doc, mansco_headers, mansco_rows, col_widths=[2.0, 4.5])
doc.add_paragraph()

# SmartForce sign block
add_heading(doc, 'SmartForce Technology Solutions — Vendor Authorization', level=2, color=DARK)
sf_rows = [
    ('Company Name',              'SmartForce Technology Solutions'),
    ('Authorized Representative', '________________________________'),
    ('Title / Position',          '________________________________'),
    ('Signature',                 '________________________________'),
    ('Date',                      '________________________________'),
]
make_table(doc, mansco_headers, sf_rows, col_widths=[2.0, 4.5])
doc.add_paragraph()

# Document control
add_heading(doc, 'Document Control', level=2, color=DARK)
dc_headers = ['Version', 'Date', 'Author', 'Change Summary']
dc_rows = [('1.0', '2026-05-20', 'SmartForce Technical PMO', 'Initial release for customer review')]
make_table(doc, dc_headers, dc_rows, col_widths=[0.7, 1.0, 2.0, 2.8])
doc.add_paragraph()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'This document is confidential and intended solely for the use of MANSCO Egypt and '
    'SmartForce Technology Solutions.\nUnauthorized distribution or reproduction is strictly prohibited.\n\n'
    'End of Document — MANSCO-SPP-MDP-001 v1.0'
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88,0x88,0x88)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = r'C:\Users\moham\OneDrive - Capture Doc\Capture Doc\SmartForce\Mansco\MANSCO Spare Parts Portal\MANSCO_Milestone_Delivery_Plan.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
